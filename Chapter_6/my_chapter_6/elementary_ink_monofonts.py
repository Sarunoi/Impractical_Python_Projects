""" Hide a message in a docx document using a white font."""
import docx
from docx.shared import RGBColor, Pt
import operator

rgb_colors_sets = {
    "red": {
        'r': 255,
        'g': 0,
        'b': 0
    },
    "green": {
        'r': 0,
        'g': 255,
        'b': 0
    },
    "blue": {
        'r': 0,
        'g': 0,
        'b': 255
    },
}


def get_paragraphs_text_list(file):
    document = docx.Document(file)
    return [paragraph.text for paragraph in document.paragraphs]


def get_template_doc_with_letterhead(template_file):
    # load template that sets style, font, margins, etc.
    template_doc = docx.Document(template_file)

    template_doc.styles['Normal'].font.name = 'Consolas'

    # add letterhead:
    template_doc.add_heading('Morland Holmes', 0)
    subtitle = template_doc.add_heading('Global Consultanting & Negotiations', 1)
    subtitle.alignment = 1
    template_doc.add_heading('', 1)
    template_doc.add_paragraph('December 17, 2015')
    template_doc.add_paragraph('')
    return template_doc


def fill_doc_with_real_msg_covering_fake_msg(doc, fake_list, real_msg):
    # for line in fake_list:
    #     if real_list and line == '':
    #         paragraph = doc.add_paragraph(real_list.pop(0))
    #         paragraph_index = len(doc.paragraphs) - 1
    #
    #         run = doc.paragraphs[paragraph_index].runs[0]
    #         run.font.color.rgb = RGBColor(
    #             **rgb_colors_sets["red"])  # make it red to test
    #
    #     else:
    #         paragraph = doc.add_paragraph(line)

    real_msg_letter_generator = get_next_letter_to_hide_generator(real_msg)()

    for line in fake_list:
        line_char_list = list(line)
        for char_position in range(len(line_char_list)):
            if line_char_list[char_position].isspace():
                try:
                    letter_to_hide = real_msg_letter_generator.__next__()
                    print(letter_to_hide)
                except StopIteration:
                    break
                else:
                    line_char_list[char_position] = letter_to_hide.upper()
        paragraph = doc.add_paragraph(''.join(line_char_list))
        for run in paragraph.runs:
            print(run.text)


def get_next_letter_to_hide_generator(msg):
    def gen():
        for letter in msg:
            yield letter
    return gen


def main():
    fake_list = get_paragraphs_text_list('fakeMessage.docx')

    real_list = [
        paragraph for paragraph in get_paragraphs_text_list('realMessage.docx')
    ]

    real_msg = ''.join(real_list).replace(' ', '')
    print(f"Real msg: {real_msg}")

    doc = get_template_doc_with_letterhead('template_challenge.docx')
    fill_doc_with_real_msg_covering_fake_msg(doc, fake_list, real_msg)

    doc.save('challenge_result.docx')

    print("Done")


if __name__ == '__main__':
    main()
